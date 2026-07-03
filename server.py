import os
import json
import time
import requests
import hashlib
import secrets
import smtplib
import re
import io
from email.mime.text import MIMEText
from typing import Optional, Dict
from bs4 import BeautifulSoup
from fastapi import FastAPI, Request, HTTPException, UploadFile, File, Form, Cookie, Depends
from fastapi.responses import JSONResponse
from fastapi.staticfiles import StaticFiles
import uvicorn
import docx
from PIL import Image
import google.generativeai as genai


app = FastAPI(title="客戶與供應商資料登錄系統 API")

# Gemini API Configuration
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "AIzaSyA3cQveOjMCWhWAOW_1k4Lvg4U1Ol8KvFs")
genai.configure(api_key=GEMINI_API_KEY)

# Ensure required directories exist
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PUBLIC_DIR = os.path.join(BASE_DIR, "public")
DATA_DIR = os.path.join(BASE_DIR, "data")
UPLOADS_DIR = os.path.join(BASE_DIR, "uploads")

for path in [PUBLIC_DIR, DATA_DIR, UPLOADS_DIR]:
    os.makedirs(path, exist_ok=True)

RECORDS_FILE = os.path.join(DATA_DIR, "records.json")
USERS_FILE = os.path.join(DATA_DIR, "users.json")
BANK_CODES_FILE = os.path.join(DATA_DIR, "bank_codes.json")

# Load bank code database
BANK_CODES = {}
if os.path.exists(BANK_CODES_FILE):
    try:
        with open(BANK_CODES_FILE, "r", encoding="utf-8") as f:
            BANK_CODES = json.load(f)
        print(f"Successfully loaded {len(BANK_CODES)} bank/branch codes.")
    except Exception as e:
        print(f"Error loading bank codes: {e}")

# Initialize records file if not exists
if not os.path.exists(RECORDS_FILE):
    with open(RECORDS_FILE, "w", encoding="utf-8") as f:
        json.dump([], f, ensure_ascii=False, indent=2)

# Initialize users file if not exists
if not os.path.exists(USERS_FILE):
    with open(USERS_FILE, "w", encoding="utf-8") as f:
        json.dump([], f, ensure_ascii=False, indent=2)

# Global in-memory stores
active_sessions: Dict[str, str] = {}
otp_store: Dict[str, dict] = {}

# Password hashing utilities
def hash_password(password: str, salt: Optional[bytes] = None) -> str:
    if not salt:
        salt = os.urandom(16)
    pw_hash = hashlib.pbkdf2_hmac('sha256', password.encode('utf-8'), salt, 100000)
    return f"{salt.hex()}:{pw_hash.hex()}"

def verify_password(password: str, stored_hash: str) -> bool:
    try:
        salt_hex, hash_hex = stored_hash.split(':')
        salt = bytes.fromhex(salt_hex)
        pw_hash = hashlib.pbkdf2_hmac('sha256', password.encode('utf-8'), salt, 100000)
        return pw_hash.hex() == hash_hex
    except Exception:
        return False

# User database loading/saving
def load_users() -> list:
    try:
        with open(USERS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []

def save_users(users: list):
    with open(USERS_FILE, "w", encoding="utf-8") as f:
        json.dump(users, f, ensure_ascii=False, indent=2)

# SMTP Config & Email sender
SMTP_SERVER = os.environ.get("SMTP_SERVER", "smtp.gmail.com")
SMTP_PORT = int(os.environ.get("SMTP_PORT", "587"))
SMTP_USER = os.environ.get("SMTP_USER", "")
SMTP_PASSWORD = os.environ.get("SMTP_PASSWORD", "")
SMTP_SENDER = os.environ.get("SMTP_SENDER", SMTP_USER or "no-reply@wavenet.com.tw")

def send_verification_email(email: str, code: str, code_type: str) -> bool:
    subject = "【潮網基本資料登錄系統】驗證碼"
    action_text = "註冊帳號" if code_type == "register" else "重設密碼"
    body = f"您好：\n\n您正在進行潮網基本資料登錄系統的{action_text}。您的 6 位數驗證碼為：\n\n  {code}\n\n驗證碼有效時間為 10 分鐘。如非本人操作，請忽略此信件。"
    
    # Always write to local latest_otp.txt file and print to terminal for dev convenience
    otp_log_path = os.path.join(DATA_DIR, "latest_otp.txt")
    with open(otp_log_path, "w", encoding="utf-8") as f:
        f.write(f"Email: {email}\nCode: {code}\nType: {code_type}\nTime: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
    print(f"\n[OTP VERIFICATION CODE FOR {email} ({code_type})]: {code}\n")
    
    if not SMTP_USER or not SMTP_PASSWORD:
        return False
        
    try:
        msg = MIMEText(body, 'plain', 'utf-8')
        msg['Subject'] = subject
        msg['From'] = SMTP_SENDER
        msg['To'] = email
        
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(SMTP_USER, SMTP_PASSWORD)
        server.sendmail(SMTP_SENDER, [email], msg.as_string())
        server.quit()
        return True
    except Exception as e:
        print(f"Error sending email via SMTP: {e}")
        return False

# Auth dependency to check session cookie
# Auth dependency - Removed to make app public
async def get_current_user():
    return "guest@public"

# Auto-calculate credit rating and limit on backend
def backend_calculate_credit(capital: float, category: str, has_debt_records: bool):
    rating = 'E2'
    limit = 0
    terms = '預收'
    
    if has_debt_records:
        rating = 'E2'
        limit = 0
        terms = '預收'
    elif category == '集團企業':
        rating = '集團企業'
        limit = 2000000
        terms = '月結 60天'
    elif category == '關係企業':
        rating = '關係企業'
        limit = 2000000
        terms = '月結 60天'
    elif category == '政府單位':
        rating = '政府單位'
        limit = 10000000
        terms = '依政府規定'
    elif category == '公股/類政府機構':
        rating = '公股/類政府機構'
        limit = 5000000
        terms = '月結 30天'
    elif category == '4A廣告公司':
        rating = 'A1'
        limit = 2000000
        terms = '月結 30天'
    elif category == '特殊分級':
        rating = 'E2'
        limit = 0
        terms = '預收'
    elif category == '診所/外國未上市公司':
        rating = 'E2'
        limit = 0
        terms = '預收'
    else:
        if capital >= 300000000:
            rating = 'A1'
            limit = 2000000
            terms = '月結 30天'
        elif capital >= 100000000:
            rating = 'A2'
            limit = 1000000
            terms = '月結 30天'
        elif capital >= 60000000:
            rating = 'B1'
            limit = 600000
            terms = '月結 30天'
        elif capital >= 30000000:
            rating = 'B2'
            limit = 400000
            terms = '月結 30天'
        elif capital >= 10000000:
            rating = 'B3'
            limit = 300000
            terms = '月結 30天'
        else:
            rating = 'E2'
            limit = 0
            terms = '預收'
            
    return rating, limit, terms

def find_bank_code_by_name(bank_name: str, branch_name: str) -> Optional[str]:
    """
    Search BANK_CODES database for a branch matching bank_name and branch_name.
    """
    if not bank_name:
        return None
        
    # Clean bank name: remove spaces, "股份有限公司", "商業銀行", "銀行"
    def clean_b_name(n):
        n = n.replace("股份有限公司", "").replace("商業銀行", "").replace("銀行", "")
        return re.sub(r'[\s\-]', '', n)
        
    # Clean branch name: remove spaces, "分行", "分公司", "營業部"
    def clean_br_name(n):
        n = n.replace("分行", "").replace("分公司", "").replace("營業部", "")
        return re.sub(r'[\s\-]', '', n)
        
    target_bank = clean_b_name(bank_name)
    target_branch = clean_br_name(branch_name) if branch_name else ""
    
    # 1. First pass: try exact match on cleaned names
    for code, info in BANK_CODES.items():
        db_bank = clean_b_name(info["bank_name"])
        db_branch = clean_br_name(info["branch_name"])
        
        if db_bank == target_bank:
            if target_branch and db_branch == target_branch:
                return code
                
    # 2. Second pass: substring match for branch if exact match fails
    if target_branch:
        for code, info in BANK_CODES.items():
            db_bank = clean_b_name(info["bank_name"])
            db_branch = clean_br_name(info["branch_name"])
            
            if db_bank == target_bank:
                # Check if target_branch is a substring of db_branch or vice versa
                if target_branch in db_branch or db_branch in target_branch:
                    return code
                    
    # 3. Third pass: return first code matching the bank if branch name is empty/unmatched
    for code, info in BANK_CODES.items():
        db_bank = clean_b_name(info["bank_name"])
        if db_bank == target_bank:
            return code
            
    return None

@app.post("/api/auth/send-code")
async def send_code(request: Request):
    data = await request.json()
    email = data.get("email", "").strip().lower()
    code_type = data.get("type", "register") # "register" or "reset"
    
    if not email:
        raise HTTPException(status_code=400, detail="請輸入 Email")
    if not (email.endswith("@wavenet.com.tw") or email.endswith(".wavenet.com.tw")):
        raise HTTPException(status_code=400, detail="請使用潮網電子郵件 (@wavenet.com.tw) 進行註冊與驗證")
        
    users = load_users()
    user_exists = any(u["email"] == email for u in users)
    
    if code_type == "register" and user_exists:
        raise HTTPException(status_code=400, detail="該信箱已註冊，請直接登入")
    elif code_type == "reset" and not user_exists:
        raise HTTPException(status_code=400, detail="該信箱尚未註冊")
        
    # Generate 6 digit OTP
    code = f"{secrets.randbelow(900000) + 100000}"
    otp_store[email] = {
        "code": code,
        "expires": time.time() + 600, # 10 mins validity
        "type": code_type
    }
    
    sent = send_verification_email(email, code, code_type)
    return {
        "success": True, 
        "message": "驗證碼已發送至您的信箱，請至收件匣或垃圾信件匣查看。" if sent else "驗證碼已發送（本地/測試模式已輸出至終端機與 data/latest_otp.txt）",
        "email_sent": sent
    }

@app.post("/api/auth/register")
async def register_user(request: Request):
    data = await request.json()
    email = data.get("email", "").strip().lower()
    password = data.get("password", "")
    code = data.get("code", "").strip()
    
    if not email or not password or not code:
        raise HTTPException(status_code=400, detail="信箱、密碼與驗證碼均為必填項")
    if not (email.endswith("@wavenet.com.tw") or email.endswith(".wavenet.com.tw")):
        raise HTTPException(status_code=400, detail="僅限潮網電子郵件 (@wavenet.com.tw) 註冊")
    if len(password) < 6:
        raise HTTPException(status_code=400, detail="密碼長度至少需為 6 個字元")
        
    otp = otp_store.get(email)
    if not otp or otp["type"] != "register" or otp["code"] != code or time.time() > otp["expires"]:
        raise HTTPException(status_code=400, detail="驗證碼錯誤或已過期，請重新發送")
        
    users = load_users()
    if any(u["email"] == email for u in users):
        raise HTTPException(status_code=400, detail="該信箱已註冊")
        
    users.append({
        "email": email,
        "password_hash": hash_password(password),
        "created_at": time.strftime("%Y-%m-%dT%H:%M:%S+08:00", time.localtime())
    })
    save_users(users)
    
    # Remove OTP after successful register
    otp_store.pop(email, None)
    return {"success": True, "message": "註冊成功，請登入！"}

@app.post("/api/auth/login")
async def login_user(request: Request):
    data = await request.json()
    email = data.get("email", "").strip().lower()
    password = data.get("password", "")
    
    if not email or not password:
        raise HTTPException(status_code=400, detail="請填寫信箱與密碼")
        
    users = load_users()
    user = next((u for u in users if u["email"] == email), None)
    
    if not user or not verify_password(password, user["password_hash"]):
        raise HTTPException(status_code=400, detail="信箱或密碼錯誤")
        
    token = secrets.token_hex(32)
    active_sessions[token] = email
    
    response = JSONResponse(content={"success": True, "message": "登入成功", "email": email})
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        max_age=86400, # 1 day
        samesite="lax",
        secure=False # set to True in production with HTTPS
    )
    return response

@app.post("/api/auth/logout")
async def logout_user(session_token: Optional[str] = Cookie(None)):
    if session_token in active_sessions:
        active_sessions.pop(session_token, None)
    response = JSONResponse(content={"success": True, "message": "已成功登出"})
    response.delete_cookie("session_token")
    return response

@app.get("/api/auth/me")
async def get_me(current_user: str = Depends(get_current_user)):
    return {"success": True, "email": current_user}

@app.post("/api/auth/reset-password")
async def reset_password(request: Request):
    data = await request.json()
    email = data.get("email", "").strip().lower()
    new_password = data.get("new_password", "")
    code = data.get("code", "").strip()
    
    if not email or not new_password or not code:
        raise HTTPException(status_code=400, detail="信箱、新密碼與驗證碼均為必填項")
    if len(new_password) < 6:
        raise HTTPException(status_code=400, detail="新密碼長度至少需為 6 個字元")
        
    otp = otp_store.get(email)
    if not otp or otp["type"] != "reset" or otp["code"] != code or time.time() > otp["expires"]:
        raise HTTPException(status_code=400, detail="驗證碼錯誤或已過期，請重新發送")
        
    users = load_users()
    user = next((u for u in users if u["email"] == email), None)
    if not user:
        raise HTTPException(status_code=400, detail="該使用者不存在")
        
    user["password_hash"] = hash_password(new_password)
    save_users(users)
    
    # Remove OTP
    otp_store.pop(email, None)
    return {"success": True, "message": "密碼重設成功，請使用新密碼登入！"}

@app.post("/api/auth/change-password")
async def change_password(request: Request, current_user: str = Depends(get_current_user)):
    data = await request.json()
    old_password = data.get("old_password", "")
    new_password = data.get("new_password", "")
    
    if not old_password or not new_password:
        raise HTTPException(status_code=400, detail="原密碼與新密碼均為必填項")
    if len(new_password) < 6:
        raise HTTPException(status_code=400, detail="新密碼長度至少需為 6 個字元")
        
    users = load_users()
    user = next((u for u in users if u["email"] == current_user), None)
    if not user:
        raise HTTPException(status_code=400, detail="該使用者不存在")
        
    if not verify_password(old_password, user["password_hash"]):
        raise HTTPException(status_code=400, detail="原密碼不正確")
        
    user["password_hash"] = hash_password(new_password)
    save_users(users)
    return {"success": True, "message": "密碼變更成功！"}


@app.get("/api/company-info")
async def get_company_info(ubn: str, current_user: str = Depends(get_current_user)):
    """
    Scrape company details from twincn.com using UBN.
    """
    if not ubn or not ubn.isdigit() or len(ubn) != 8:
        raise HTTPException(status_code=400, detail="請輸入正確的 8 位數統一編號")
        
    url = f"https://www.twincn.com/item.aspx?no={ubn}"
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8',
        'Accept-Language': 'zh-TW,zh;q=0.9,en-US;q=0.8,en;q=0.7',
    }
    
    try:
        response = requests.get(url, headers=headers, timeout=10)
        # Check if redirected to the homepage (means UBN not found or invalid)
        if response.history:
            # Check if any redirect in history led to twincn.com/ (home page)
            final_url = response.url.rstrip('/')
            if final_url == 'https://www.twincn.com' or final_url == 'http://www.twincn.com':
                return {"success": False, "message": "查無此統一編號之公司資料（已被重新導向首頁）"}
        
        if response.status_code != 200:
            return {"success": False, "message": f"無法存取該網頁，錯誤碼：{response.status_code}"}
            
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Check if the page is just the home page
        title_text = soup.title.get_text() if soup.title else ""
        if "台灣公司網" == title_text.strip() and not soup.find('td', text=re.compile("統一編號|公司名稱")):
            # Double check if company data td is absent
            # Check if there is any td containing the UBN
            td_ubn = soup.find('td', text=ubn)
            if not td_ubn:
                return {"success": False, "message": "查無此統一編號之公司資料"}
        
        # Parse fields from the TDs
        company_name = ""
        address = ""
        capital_str = ""
        phone = ""
        
        tds = soup.find_all('td')
        for i in range(len(tds) - 1):
            text = tds[i].get_text(" ", strip=True)
            next_text = tds[i+1].get_text(" ", strip=True)
            
            if text == "公司名稱" or text == "公司全名":
                company_name = next_text
            elif text == "公司所在地":
                address = next_text
            elif "資本總額(元)" in text or "資本額" in text:
                if not capital_str:
                    capital_str = next_text
            elif text == "電話":
                if not phone:
                    phone = next_text
        
        # Fallback for company name: title of the page sometimes contains it
        if not company_name and soup.title:
            title_parts = soup.title.get_text().split('·')
            if len(title_parts) > 0 and title_parts[0] != "台灣公司網":
                company_name = title_parts[0].strip()
                
        # Clean phone number (remove (來源:...) annotations)
        if phone:
            clean_phone = re.sub(r'\(來源:[^\)]+\)', '', phone)
            # Split and keep unique tokens
            tokens = [t.strip() for t in clean_phone.split() if t.strip()]
            unique_tokens = []
            for t in tokens:
                if t not in unique_tokens:
                    unique_tokens.append(t)
            phone = ", ".join(unique_tokens)
            
        # Parse numeric capital
        capital_val = 0
        if capital_str:
            clean_cap = re.sub(r'[^\d]', '', capital_str)
            if clean_cap:
                capital_val = int(clean_cap)
        
        if not company_name and not address:
            return {"success": False, "message": "解析失敗，無法在此頁面找到公司基本資料欄位"}
            
        return {
            "success": True,
            "company_name": company_name,
            "address": address,
            "capital": capital_val,
            "capital_str": capital_str,
            "phone": phone
        }
        
    except Exception as e:
        return {"success": False, "message": f"發生錯誤：{str(e)}"}

@app.get("/api/records")
async def get_records(current_user: str = Depends(get_current_user)):
    """
    Get all registered customer and supplier records.
    """
    try:
        with open(RECORDS_FILE, "r", encoding="utf-8") as f:
            records = json.load(f)
        return records
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"讀取資料庫失敗：{str(e)}")

@app.post("/api/register")
async def register_info(request: Request, current_user: str = Depends(get_current_user)):
    """
    Register customer or supplier details, save files, and record to JSON database.
    """
    try:
        form_data = await request.form()
        
        register_type = form_data.get("register_type")
        if not register_type or register_type not in ["customer", "supplier"]:
            return JSONResponse(status_code=400, content={"success": False, "message": "註冊類別不正確"})
            
        # Extract fields
        record = {
            "id": f"rec_{int(time.time())}",
            "timestamp": time.strftime("%Y-%m-%dT%H:%M:%S+08:00", time.localtime()),
            "register_type": register_type
        }
        
        # Populate all text fields dynamically
        for key in form_data.keys():
            # Skip file inputs as they are handled separately
            if hasattr(form_data[key], "filename"):
                continue
            record[key] = form_data.get(key)
            
        # Overwrite/calculate credit rating fields for non-Caroline
        if register_type == "customer":
            if current_user != "caroline@wavenet.com.tw":
                # Force category and debt defaults, and recalculate limits
                capital = float(record.get("capital", 0))
                category = "一般公司"
                has_debt_records = False
                rating, limit, terms = backend_calculate_credit(capital, category, has_debt_records)
                
                record["cust_category"] = category
                record["has_debt_records"] = "false"
                record["credit_limit"] = str(limit)
                record["credit_rating"] = rating
                record["payment_terms"] = terms
            
        # Save directory path checking and creation
        os.makedirs(UPLOADS_DIR, exist_ok=True)
        
        # Helper to save files
        async def save_uploaded_file(field_name, file_label):
            file_obj = form_data.get(field_name)
            if file_obj and hasattr(file_obj, "filename") and file_obj.filename:
                # Get extension
                _, ext = os.path.splitext(file_obj.filename)
                if not ext:
                    ext = ".jpg"  # Default fallback
                # Make a secure filename
                ubn_id = form_data.get("ubn") or form_data.get("ubn_or_id") or "unknown"
                clean_ubn = re.sub(r'[^\w\-]', '_', ubn_id)
                safe_name = f"{record['id']}_{clean_ubn}_{file_label}{ext}"
                file_path = os.path.join(UPLOADS_DIR, safe_name)
                
                content = await file_obj.read()
                with open(file_path, "wb") as f:
                    f.write(content)
                return f"/uploads/{safe_name}"
            return None
            
        # Save files based on type
        record["passbook_url"] = await save_uploaded_file("passbook_file", "passbook")
        record["id_front_url"] = await save_uploaded_file("id_front_file", "id_front")
        record["id_back_url"] = await save_uploaded_file("id_back_file", "id_back")
        
        # Read current database
        with open(RECORDS_FILE, "r", encoding="utf-8") as f:
            records = json.load(f)
            
        records.insert(0, record)  # Insert newest at top
        
        # Write back to database
        with open(RECORDS_FILE, "w", encoding="utf-8") as f:
            json.dump(records, f, ensure_ascii=False, indent=2)
            
        # Generate and fill docx for customers
        download_url = None
        if register_type == "customer":
            try:
                template_path = os.path.join(BASE_DIR, "W27-客戶基本資料暨信用額度申請表.docx")
                clean_ubn = re.sub(r'[^\w\-]', '_', record.get("ubn", "unknown"))
                doc_filename = f"{record['id']}_{clean_ubn}_customer_application.docx"
                doc_filepath = os.path.join(UPLOADS_DIR, doc_filename)
                
                if os.path.exists(template_path):
                    doc = docx.Document(template_path)
                    
                    # 1. Fill paragraphs (Apply type and date)
                    year = time.strftime("%Y")
                    month = time.strftime("%m")
                    day = time.strftime("%d")
                    
                    for p in doc.paragraphs:
                        if "申請類別" in p.text:
                            p.text = p.text.replace("□新增客戶", "■新增客戶" if record.get("apply_type") == "新增" else "□新增客戶")
                            p.text = p.text.replace("□變更資料", "■變更資料" if record.get("apply_type") == "變更" else "□變更資料")
                            p.text = p.text.replace("□停用", "■停用" if record.get("apply_type") == "停用" else "□停用")
                            p.text = p.text.replace("年", f"{year} 年")
                            p.text = p.text.replace("月", f"{month} 月")
                            p.text = p.text.replace("日", f"{day} 日")
                            
                    # Helpers for filling cells
                    def fill_cell_text(cell, text):
                        if len(cell.paragraphs) > 0:
                            cell.paragraphs[0].text = str(text)
                            for ep in cell.paragraphs[1:]:
                                ep.text = ""
                        else:
                            cell.text = str(text)
                            
                    def replace_cell_checkboxes(cell, old_val, new_val):
                        for cp in cell.paragraphs:
                            if old_val in cp.text:
                                cp.text = cp.text.replace(old_val, new_val)
                                
                    table = doc.tables[0]
                    
                    # Row 0: Apply Dept and User
                    fill_cell_text(table.rows[0].cells[1], record.get("apply_dept", ""))
                    fill_cell_text(table.rows[0].cells[4], record.get("apply_user", ""))
                    
                    # Row 2: Company Name (發票抬頭)
                    fill_cell_text(table.rows[2].cells[1], record.get("company_name", ""))
                    
                    # Row 3: UBN
                    fill_cell_text(table.rows[3].cells[1], record.get("ubn", ""))
                    
                    # Row 4: Identity Type
                    ident = record.get("identity_type", "")
                    cell_ident = table.rows[4].cells[1]
                    if ident == "國內公司":
                        replace_cell_checkboxes(cell_ident, "□ 國內公司", "■ 國內公司")
                    elif ident == "國外公司":
                        replace_cell_checkboxes(cell_ident, "□ 國外公司", "■ 國外公司")
                    elif ident == "個人":
                        replace_cell_checkboxes(cell_ident, "□ 個人", "■ 個人")
                    else:
                        replace_cell_checkboxes(cell_ident, "□其他:___________________________", f"■其他: {ident}")
                        
                    # Row 5: Invoice Type
                    inv = record.get("invoice_type", "")
                    cell_inv = table.rows[5].cells[1]
                    if "統一發票" in inv:
                        replace_cell_checkboxes(cell_inv, "□開立統一發票", "■開立統一發票")
                    elif "Invoice" in inv:
                        replace_cell_checkboxes(cell_inv, "□ 開立Invoice", "■ 開立Invoice")
                        
                    # Row 7: Registered Address
                    fill_cell_text(table.rows[7].cells[1], record.get("registered_address", ""))
                    
                    # Row 8: Mailing Address
                    fill_cell_text(table.rows[8].cells[1], record.get("mailing_address", ""))
                    
                    # Row 9: Phone
                    fill_cell_text(table.rows[9].cells[1], record.get("phone", ""))
                    
                    # Row 10: Contact Person
                    fill_cell_text(table.rows[10].cells[1], record.get("contact_person", ""))
                    
                    # Row 11: Email
                    fill_cell_text(table.rows[11].cells[1], record.get("email", ""))
                    
                    # Row 13: Payment Terms
                    terms = record.get("payment_terms", "")
                    cell_terms = table.rows[13].cells[1]
                    if "30天" in terms:
                        replace_cell_checkboxes(cell_terms, "□ 30天", "■ 30天")
                    elif "60天" in terms:
                        replace_cell_checkboxes(cell_terms, "□ 60天", "■ 60天")
                    elif "90天" in terms:
                        replace_cell_checkboxes(cell_terms, "□90天", "■90天")
                    elif "即期" in terms:
                        replace_cell_checkboxes(cell_terms, "□即期", "■即期")
                    elif "預收" in terms:
                        replace_cell_checkboxes(cell_terms, "□預收", "■預收")
                    else:
                        replace_cell_checkboxes(cell_terms, "□其他:_____________________", f"■其他: {terms}")
                        
                    # Row 14: Payment Method
                    pm = record.get("payment_method", "")
                    cell_pm = table.rows[14].cells[1]
                    if "電匯" in pm:
                        replace_cell_checkboxes(cell_pm, "□電匯", "■電匯")
                    elif "票據" in pm:
                        replace_cell_checkboxes(cell_pm, "□票據", "■票據")
                        
                    # Row 15: Currency
                    curr = record.get("currency", "")
                    cell_curr = table.rows[15].cells[1]
                    if "TWD" in curr:
                        replace_cell_checkboxes(cell_curr, "□TWD", "■TWD")
                    elif "USD" in curr:
                        replace_cell_checkboxes(cell_curr, "□ USD", "■ USD")
                    elif "RMB" in curr:
                        replace_cell_checkboxes(cell_curr, "□ RMB", "■ RMB")
                    elif "JPY" in curr:
                        replace_cell_checkboxes(cell_curr, "□  JPY", "■  JPY")
                    else:
                        replace_cell_checkboxes(cell_curr, "□ 其他:_____________", f"■ 其他: {curr}")
                        
                    # Row 16: Notes
                    fill_cell_text(table.rows[16].cells[1], record.get("notes", ""))
                    
                    # Row 17: Credit Rating & Credit Limit
                    fill_cell_text(table.rows[17].cells[1], record.get("credit_rating", ""))
                    try:
                        limit_val = float(record.get("credit_limit", 0))
                        limit_str = f"{limit_val:,.0f}"
                    except Exception:
                        limit_str = str(record.get("credit_limit", 0))
                    fill_cell_text(table.rows[17].cells[5], limit_str)
                    
                    # Row 18: Debt Records
                    debt = record.get("has_debt_records", "false")
                    cell_debt = table.rows[18].cells[0]
                    if debt == "true":
                        replace_cell_checkboxes(cell_debt, "□有", "■有")
                    else:
                        replace_cell_checkboxes(cell_debt, "□無", "■無")
                        
                    doc.save(doc_filepath)
                    download_url = f"/uploads/{doc_filename}"
            except Exception as docx_err:
                print(f"Error generating filled docx: {docx_err}")
        elif register_type == "supplier":
            try:
                template_path = os.path.join(BASE_DIR, "W29-供應商資料表.docx")
                clean_ubn = re.sub(r'[^\w\-]', '_', record.get("ubn_or_id", "unknown"))
                doc_filename = f"{record['id']}_{clean_ubn}_supplier_application.docx"
                doc_filepath = os.path.join(UPLOADS_DIR, doc_filename)
                
                if os.path.exists(template_path):
                    doc = docx.Document(template_path)
                    
                    # 1. Fill paragraphs (Apply type and date)
                    year = time.strftime("%Y")
                    month = time.strftime("%m")
                    day = time.strftime("%d")
                    
                    for p in doc.paragraphs:
                        if "申請類別" in p.text:
                            p.text = p.text.replace("□新增", "■新增" if record.get("apply_type") == "新增" else "□新增")
                            p.text = p.text.replace("□變更資料", "■變更資料" if record.get("apply_type") == "變更" else "□變更資料")
                            p.text = p.text.replace("□停用", "■停用" if record.get("apply_type") == "停用" else "□停用")
                            p.text = p.text.replace("年", f"{year} 年")
                            p.text = p.text.replace("月", f"{month} 月")
                            p.text = p.text.replace("日", f"{day} 日")
                            
                    # Helpers for filling cells
                    def fill_cell_text(cell, text):
                        if len(cell.paragraphs) > 0:
                            cell.paragraphs[0].text = str(text)
                            for ep in cell.paragraphs[1:]:
                                ep.text = ""
                        else:
                            cell.text = str(text)
                            
                    def replace_cell_checkboxes(cell, old_val, new_val):
                        for cp in cell.paragraphs:
                            if old_val in cp.text:
                                cp.text = cp.text.replace(old_val, new_val)
                                
                    table = doc.tables[0]
                    
                    # Row 1: Company Name (發票抬頭)
                    fill_cell_text(table.rows[1].cells[1], record.get("company_name", ""))
                    
                    # Row 2: UBN / ID
                    fill_cell_text(table.rows[2].cells[1], record.get("ubn_or_id", ""))
                    
                    # Row 3: Identity Type
                    ident = record.get("identity_type", "")
                    cell_ident = table.rows[3].cells[1]
                    if ident == "國內公司":
                        replace_cell_checkboxes(cell_ident, "□ 國內公司", "■ 國內公司")
                    elif ident == "國外公司":
                        replace_cell_checkboxes(cell_ident, "□ 國外公司", "■ 國外公司")
                    elif ident == "個人":
                        replace_cell_checkboxes(cell_ident, "□ 個人", "■ 個人")
                    else:
                        replace_cell_checkboxes(cell_ident, "□其他:___________________________", f"■其他: {ident}")
                        
                    # Row 4: Invoice Type
                    inv = record.get("invoice_type", "")
                    cell_inv = table.rows[4].cells[1]
                    if "三聯式" in inv:
                        replace_cell_checkboxes(cell_inv, "□開立三聯式發票", "■開立三聯式發票")
                        replace_cell_checkboxes(cell_inv, "□ 開立三聯式發票", "■ 開立三聯式發票")
                    elif "電子" in inv:
                        replace_cell_checkboxes(cell_inv, "□開立電子發票", "■開立電子發票")
                        replace_cell_checkboxes(cell_inv, "□ 開立電子發票", "■ 開立電子發票")
                    elif "收據" in inv:
                        replace_cell_checkboxes(cell_inv, "□開立收據/勞報單", "■開立收據/勞報單")
                        replace_cell_checkboxes(cell_inv, "□ 開立收據/勞報單", "■ 開立收據/勞報單")
                        replace_cell_checkboxes(cell_inv, "□開立收據", "■開立收據")
                        replace_cell_checkboxes(cell_inv, "□ 開立收據", "■ 開立收據")
                    elif "Invoice" in inv:
                        replace_cell_checkboxes(cell_inv, "□開立Invoice", "■開立Invoice")
                        replace_cell_checkboxes(cell_inv, "□ 開立Invoice", "■ 開立Invoice")
                        
                    # Row 5: Registered Address
                    fill_cell_text(table.rows[5].cells[1], record.get("registered_address", ""))
                    
                    # Row 6: Mailing Address
                    fill_cell_text(table.rows[6].cells[1], record.get("mailing_address", ""))
                    
                    # Row 7: Phone (Col 1) and Contact Person (Col 5)
                    fill_cell_text(table.rows[7].cells[1], record.get("phone", ""))
                    fill_cell_text(table.rows[7].cells[5], record.get("contact_person", ""))
                    
                    # Row 8: Email
                    fill_cell_text(table.rows[8].cells[1], record.get("email", ""))
                    
                    # Row 10: Payment Method
                    pm = record.get("payment_method", "")
                    cell_pm = table.rows[10].cells[1]
                    if "電匯" in pm:
                        replace_cell_checkboxes(cell_pm, "□電匯", "■電匯")
                        replace_cell_checkboxes(cell_pm, "□ 電匯", "■ 電匯")
                    elif "信用卡" in pm:
                        replace_cell_checkboxes(cell_pm, "□信用卡", "■信用卡")
                        replace_cell_checkboxes(cell_pm, "□ 信用卡", "■ 信用卡")
                        
                    # Row 11: Currency
                    curr = record.get("currency", "")
                    cell_curr = table.rows[11].cells[1]
                    if "TWD" in curr:
                        replace_cell_checkboxes(cell_curr, "□TWD", "■TWD")
                        replace_cell_checkboxes(cell_curr, "□ TWD", "■ TWD")
                    elif "USD" in curr:
                        replace_cell_checkboxes(cell_curr, "□ USD", "■ USD")
                        replace_cell_checkboxes(cell_curr, "□  USD", "■  USD")
                    elif "RMB" in curr:
                        replace_cell_checkboxes(cell_curr, "□ RMB", "■ RMB")
                        replace_cell_checkboxes(cell_curr, "□  RMB", "■  RMB")
                    elif "JPY" in curr:
                        replace_cell_checkboxes(cell_curr, "□  JPY", "■  JPY")
                    else:
                        replace_cell_checkboxes(cell_curr, "□ 其他:_____________", f"■ 其他: {curr}")
                        
                    # Row 12: Bank Name (Col 1) and Branch Name (Col 6)
                    fill_cell_text(table.rows[12].cells[1], record.get("bank_name", ""))
                    fill_cell_text(table.rows[12].cells[6], record.get("branch_name", ""))
                    
                    # Row 13: Bank Code
                    fill_cell_text(table.rows[13].cells[1], record.get("bank_code", ""))
                    
                    # Row 14: Bank Account Name
                    fill_cell_text(table.rows[14].cells[1], record.get("bank_account_name", ""))
                    
                    # Row 15: Bank Account Number
                    fill_cell_text(table.rows[15].cells[1], record.get("bank_account_number", ""))
                    
                    # Row 16: Notes
                    fill_cell_text(table.rows[16].cells[1], record.get("notes", ""))
                    
                    # Row 18: Payment Terms
                    terms = record.get("payment_terms", "")
                    cell_terms = table.rows[18].cells[1]
                    if "30天" in terms:
                        replace_cell_checkboxes(cell_terms, "□ 30天", "■ 30天")
                        replace_cell_checkboxes(cell_terms, "□30天", "■30天")
                    elif "60天" in terms:
                        replace_cell_checkboxes(cell_terms, "□ 60天", "■ 60天")
                        replace_cell_checkboxes(cell_terms, "□60天", "■60天")
                    elif "90天" in terms:
                        replace_cell_checkboxes(cell_terms, "□ 90天", "■ 90天")
                        replace_cell_checkboxes(cell_terms, "□90天", "■90天")
                    else:
                        replace_cell_checkboxes(cell_terms, "□其他:_____________________", f"■其他: {terms}")
                        
                    reason = record.get("payment_terms_reason", "")
                    if reason:
                        replace_cell_checkboxes(cell_terms, "請說明原因:", f"請說明原因: {reason}")
                        
                    # Row 20: Apply User
                    fill_cell_text(table.rows[20].cells[3], record.get("apply_user", ""))
                    
                    doc.save(doc_filepath)
                    download_url = f"/uploads/{doc_filename}"
            except Exception as docx_err:
                print(f"Error generating filled supplier docx: {docx_err}")
            
        return {"success": True, "message": "資料登錄成功", "record_id": record["id"], "download_url": download_url}
        
    except Exception as e:
        return JSONResponse(status_code=500, content={"success": False, "message": f"儲存失敗：{str(e)}"})

@app.post("/api/ocr/recognize")
async def ocr_recognize(
    file: UploadFile = File(...),
    doc_type: str = Form(...),
    current_user: str = Depends(get_current_user)
):
    """
    Recognize bank passbook or ID card image and extract fields as JSON.
    """
    if doc_type not in ["passbook", "id_front", "id_back"]:
        raise HTTPException(status_code=400, detail="不支援的證件類型")
        
    try:
        # Read file bytes
        file_bytes = await file.read()
        image = Image.open(io.BytesIO(file_bytes))
        
        # Prepare prompts
        if doc_type == "passbook":
            prompt = """
            You are an OCR assistant for a Taiwanese banking system. Analyze this bank passbook cover image and extract:
            1. bank_name: The name of the bank in Traditional Chinese (e.g., "兆豐國際商業銀行" or "國泰世華商業銀行")
            2. branch_name: The name of the branch in Traditional Chinese (e.g., "金控總部分行" or "建國分行")
            3. bank_code: A 7-digit code combining the 3-digit bank code and 4-digit branch code (e.g., "0170009" or "0130152"). If you only find the 3-digit bank code, append "0000" or try to infer it. Mega bank is 017. Cathay United is 013.
            4. bank_account_name: The account holder's name in Traditional Chinese (e.g., "潮網科技股份有限公司" or "張三").
            5. bank_account_number: The bank account number containing only digits (remove hyphens, spaces, e.g., "01709123456").
            
            Output your response ONLY as a JSON object matching this schema:
            {
              "bank_name": "...",
              "branch_name": "...",
              "bank_code": "...",
              "bank_account_name": "...",
              "bank_account_number": "..."
            }
            """
        elif doc_type == "id_front":
            prompt = """
            You are an OCR assistant for a Taiwanese identity verification system. Analyze this Taiwan ID Card (正面) image and extract:
            1. id_number: The National Identification Number (身分證字號) starting with an English letter followed by 9 digits (e.g., "A123456789").
            2. name: The individual's full name in Traditional Chinese (e.g., "張三").
            3. birthday: The birthday in Republic of China (ROC) calendar format (e.g., "民國78年10月10日").
            
            Output your response ONLY as a JSON object matching this schema:
            {
              "id_number": "...",
              "name": "...",
              "birthday": "..."
            }
            """
        else: # id_back
            prompt = """
            You are an OCR assistant for a Taiwanese identity verification system. Analyze this Taiwan ID Card (反面) image and extract:
            1. address: The individual's registered address (住址) in Traditional Chinese (e.g., "台北市大安區新生南路三段123號").
            
            Output your response ONLY as a JSON object matching this schema:
            {
              "address": "..."
            }
            """
            
        # Call Gemini model
        model = genai.GenerativeModel('gemini-2.5-flash')
        generation_config = {"response_mime_type": "application/json"}
        
        response = model.generate_content(
            [prompt, image],
            generation_config=generation_config
        )
        
        # Parse JSON to verify correctness
        data = json.loads(response.text)
        
        # Clean/refine bank data using our bank_codes database and perform reverse lookup
        if doc_type == "passbook":
            b_name = data.get("bank_name", "").strip()
            br_name = data.get("branch_name", "").strip()
            b_code = data.get("bank_code", "").strip()
            
            # Clean non-digits from code
            b_code = re.sub(r'[^\d]', '', b_code)
            
            matched_code = None
            if len(b_code) == 7 and b_code in BANK_CODES:
                matched_code = b_code
            else:
                # Code is empty or invalid. Attempt reverse lookup using names
                matched_code = find_bank_code_by_name(b_name, br_name)
                
            if matched_code:
                canonical = BANK_CODES[matched_code]
                data["bank_name"] = canonical["bank_name"]
                data["branch_name"] = canonical["branch_name"]
                data["bank_code"] = canonical["bank_code"]
            else:
                data["bank_code"] = b_code
                
        return {"success": True, "data": data}
        
    except json.JSONDecodeError as je:
        print(f"Gemini did not return valid JSON: {response.text}")
        return JSONResponse(status_code=500, content={"success": False, "message": "AI 解析結果格式不正確"})
    except Exception as e:
        print(f"OCR recognition error: {e}")
        return JSONResponse(status_code=500, content={"success": False, "message": f"AI 辨識失敗：{str(e)}"})

@app.get("/api/bank/lookup")
async def bank_lookup(code: str, current_user: str = Depends(get_current_user)):
    """
    Look up bank details using 7-digit branch code from the database.
    """
    code = code.strip()
    if code in BANK_CODES:
        return {"success": True, "data": BANK_CODES[code]}
    return JSONResponse(
        status_code=404, 
        content={"success": False, "message": "查無此銀行與分行代碼，請確認是否輸入正確。"}
    )

# Serve uploads directory
app.mount("/uploads", StaticFiles(directory=UPLOADS_DIR), name="uploads")

# Serve public directory
app.mount("/", StaticFiles(directory=PUBLIC_DIR, html=True), name="public")

if __name__ == "__main__":
    port = int(os.getenv("PORT", 8080))
    print(f"啟動客戶與供應商資料登錄系統，請在瀏覽器打開 http://localhost:{port}")
    uvicorn.run("server:app", host="0.0.0.0", port=port, reload=False)
